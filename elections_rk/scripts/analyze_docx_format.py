#!/usr/bin/env python3
"""
ЭТАП 0: Анализ формата DOCX для извлечения границ участков.
Выводит сырой текст для принятия архитектурных решений.
"""

from docx import Document
from pathlib import Path
import sys

BOUNDARIES_DIR = Path(__file__).parent.parent / "data" / "boundaries"


def analyze_docx(docx_path: Path, max_paragraphs: int = 150):
    """Извлечь и вывести параграфы из DOCX."""
    print("=" * 80)
    print(f"ФАЙЛ: {docx_path.name}")
    print(f"РАЗМЕР: {docx_path.stat().st_size:,} байт")
    print("=" * 80)
    
    doc = Document(docx_path)
    total = len(doc.paragraphs)
    print(f"ВСЕГО ПАРАГРАФОВ: {total}\n")
    
    print("-" * 80)
    print("ПЕРВЫЕ ПАРАГРАФЫ (для анализа структуры):")
    print("-" * 80)
    
    for i, para in enumerate(doc.paragraphs[:max_paragraphs]):
        text = para.text.strip()
        if text:
            # Обрезаем длинные строки
            display = text[:140] + "..." if len(text) > 140 else text
            print(f"[{i:03d}] {display}")
    
    print("\n" + "=" * 80)
    print("КОНЕЦ ВЫБОРКИ")
    print("=" * 80)


def main():
    # Находим самый большой DOCX (вероятно, содержит много участков)
    docx_files = list(BOUNDARIES_DIR.glob("*.docx"))
    
    if not docx_files:
        print("НЕТ DOCX ФАЙЛОВ В", BOUNDARIES_DIR)
        sys.exit(1)
    
    # Сортируем по размеру (убывание)
    docx_files.sort(key=lambda p: p.stat().st_size, reverse=True)
    
    print("НАЙДЕННЫЕ DOCX ФАЙЛЫ:")
    for f in docx_files[:5]:
        print(f"  {f.stat().st_size:>10,} байт  {f.name[:60]}")
    print()
    
    # Анализируем самый большой
    analyze_docx(docx_files[0])


if __name__ == "__main__":
    main()
